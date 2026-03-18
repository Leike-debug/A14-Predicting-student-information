import os
import glob
import pandas as pd
from docx import Document
import re
import csv
from io import StringIO

# 屏蔽 Pandas 内部警告
pd.options.mode.chained_assignment = None

# ================== 配置路径 ==================
DOC_DIR = r"D:\github\A14-Predicting-student-information\data\doc type"
DATA_DIR = r"D:\github\A14-Predicting-student-information\data\excel type"
OUTPUT_DIR = r"D:\github\A14-Predicting-student-information\data\translated"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# ================== 阶段 1：Word 智能提取 (表内嵌字典 + 纯文本字典) ==================
def process_text_line(cells, column_map, global_dict):
    if len(cells) < 2: return
    col_en = cells[0].strip().upper()
    col_cn_raw = cells[-1].strip() 
    
    dict_matches = re.findall(r'(\d+)\s*[:：\-\.,，]?\s*([\u4e00-\u9fa5]+)', col_cn_raw)
    
    match = re.match(r'^[\u4e00-\u9fa5]+', col_cn_raw.replace(' ', ''))
    col_cn_pure = match.group(0) if match else col_en
    col_cn_pure = re.sub(r'[\(（\[【].*?[\)）\]】]', '', col_cn_pure).strip(' ,:;，：；')

    if re.match(r'^[A-Z0-9_\-]+$', col_en) and col_cn_pure and col_cn_pure != col_en:
        column_map[col_en] = col_cn_pure
        
    if dict_matches:
        base_key = col_cn_pure.replace("代码", "")
        if base_key not in global_dict:
            global_dict[base_key] = {}
        for code, meaning in dict_matches:
            if meaning not in ['和', '与', '或', '及']: 
                global_dict[base_key][code] = meaning

def parse_word_metadata(doc_dir):
    print("[阶段1] 启动 Word 解析 (提取表头映射 & 显性文字代号字典)...")
    global_dict = {}  
    column_map = {}   
    
    doc_files = glob.glob(os.path.join(doc_dir, "*.docx"))
    if not doc_files: return global_dict, column_map

    for file_path in doc_files:
        try:
            doc = Document(file_path)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    process_text_line(cells, column_map, global_dict)

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text: continue
                
                if ',' in text and not text.endswith(':'):
                    try:
                        cells = next(csv.reader(StringIO(text)))
                        process_text_line(cells, column_map, global_dict)
                    except:
                        pass
                else:
                    m1 = re.match(r'^([A-Za-z0-9_\-]+)\s*[：:\-\.]\s*(.+)$', text)
                    if m1:
                        global_dict["通用"] = global_dict.get("通用", {})
                        global_dict["通用"][m1.group(1).strip()] = m1.group(2).strip()
        except Exception:
            pass

    global_dict = {k: v for k, v in global_dict.items() if len(v) > 0}
    print(f"  -> Word提取完毕：获取 {len(column_map)} 个表头映射，以及 {len(global_dict)} 个显性文字字典簇。")
    return global_dict, column_map


# ================== 阶段 2：Excel 逆向自学习 ==================
def clean_column_names(df, column_map):
    new_cols = []
    for col in df.columns:
        col_str = str(col).strip()
        col_upper = col_str.upper()
        if col_upper in column_map:
            new_cols.append(column_map[col_upper])
        elif re.search(r'[\(（](.*?)[\)）]', col_str):
            new_cols.append(re.search(r'[\(（](.*?)[\)）]', col_str).group(1).strip())
        else:
            new_cols.append(col_str)
    df.columns = new_cols
    return df

def enrich_dict_from_data(data_dir, global_dict, column_map):
    print("\n[阶段2] 启动 Excel 逆向自学习 (找回表格中的隐性代号字典)...")
    files = [f for f in os.listdir(data_dir) if re.match(r'.*\.(csv|xlsx|xls)$', f, re.IGNORECASE)]
    
    for file_name in files:
        file_path = os.path.join(data_dir, file_name)
        try:
            if file_name.lower().endswith(".csv"):
                df = pd.read_csv(file_path, nrows=3000, low_memory=False, dtype=str, on_bad_lines='skip')
            else:
                df = pd.read_excel(file_path, nrows=3000, dtype=str)

            df = clean_column_names(df, column_map)
            # 🔥 防御装甲：去重列名，防止后续出现 DataFrame object has no attribute 'str' 报错
            df = df.loc[:, ~df.columns.duplicated()]
            cols = df.columns.tolist()

            for col in cols:
                if col.endswith("名称") or col.endswith("含义"):
                    base_name = col.replace("名称", "").replace("含义", "")
                    code_col = None
                    if base_name + "代码" in cols: code_col = base_name + "代码"
                    elif base_name in cols: code_col = base_name 
                        
                    if code_col:
                        mapping_df = df[[code_col, col]].dropna().drop_duplicates()
                        base_key = base_name.replace("代码", "")
                        if base_key not in global_dict: global_dict[base_key] = {}
                        for _, row in mapping_df.iterrows():
                            code = re.sub(r'\.0$', '', str(row[code_col]).strip())
                            meaning = str(row[col]).strip()
                            if code and meaning and code.lower() not in ['nan', 'none']:
                                global_dict[base_key][code] = meaning
        except Exception:
            pass
    print(f"  -> 自学习完毕！代号知识库扩充至 {len(global_dict)} 个。")
    return global_dict


# ================== 阶段 3：强制 ID 归一化与 CSV 翻译输出 ==================
def identify_id_type(df):
    id_map = {}
    
    # 🔥 终极补丁：不管是正规的 CARDID，还是拼写错误的 CARDLD，全部一网打尽！
    primary_kw = ['学号', 'XH', '学生编号', 'STUDENT', 'XUEHAO', 'CARDID', 'CARDLD']
    # 替补学号库
    fallback_kw = ['发布人账号', '账号', '登录名', '回帖人账号', 'USERID', 'USERNAME', '人员编号']
    
    for col in df.columns:
        name = str(col).upper() # 转大写匹配
        
        # 拉黑名单：防误伤老师的账号
        if any(x in name for x in ['TEACHER', '老师', '辅导员', 'REVIEW', '批阅', 'TARGET']): 
            continue
        
        # 优先级判断
        if any(k in name for k in primary_kw):
            if "student_id" not in id_map.values():
                id_map[col] = "student_id"
        elif any(k in name for k in fallback_kw) and "student_id" not in id_map.values():
            id_map[col] = "student_id"
            
    return id_map

def normalize_id_series(series):
    series = series.astype(str).str.strip().str.replace(r'\.0$', '', regex=True)
    series = series.replace(["nan", "None", "", "<NA>", "null"], pd.NA)
    valid_mask = series.notna()
    lens = series[valid_mask].str.len()
    if len(lens) > 0:
        series.loc[valid_mask] = series.loc[valid_mask].str.zfill(int(lens.mode()[0]))
    return series

def translate_data_files(data_dir, output_dir, global_dict, column_map):
    print("\n[阶段3] 开始全量数据清洗、ID强制归一化与代号翻译...")
    files = [f for f in os.listdir(data_dir) if re.match(r'.*\.(csv|xlsx|xls)$', f, re.IGNORECASE)]

    for file_name in files:
        file_path = os.path.join(data_dir, file_name)
        try:
            if file_name.lower().endswith(".csv"):
                df = pd.read_csv(file_path, low_memory=False, dtype=str, on_bad_lines='skip')
            else:
                df = pd.read_excel(file_path, dtype=str)

            # 1. 洗表头
            df = clean_column_names(df, column_map)
            
            # 🔥 核心防爆补丁：剔除名字完全重复的列（保留第一列），杜绝合并时内存溢出和属性报错
            df = df.loc[:, ~df.columns.duplicated()]

            # 2. 学号统一 (强制重命名)
            id_map = identify_id_type(df)
            if "student_id" in id_map.values():
                student_cols = [k for k, v in id_map.items() if v == "student_id"]
                if student_cols:
                    old_student_col = student_cols[0]
                    if old_student_col != '学号_统一':
                        if '学号_统一' not in df.columns:
                            df.rename(columns={old_student_col: '学号_统一'}, inplace=True)
                            id_map['学号_统一'] = id_map.pop(old_student_col)
                        else:
                            id_map.pop(old_student_col, None)
                    if '学号_统一' in df.columns:
                        df['学号_统一'] = normalize_id_series(df['学号_统一'])

            for col, id_type in id_map.items():
                if col != '学号_统一' and col in df.columns:
                    df[col] = normalize_id_series(df[col])

            # 3. 代号含义翻译
            translated_cols = 0
            for col in df.columns.tolist():
                if "含义" in col or "名称" in col: continue
                
                for dict_key, mapping in global_dict.items():
                    if dict_key in col or col in dict_key:
                        # 此时 df[col] 绝对安全，必是单列 Series
                        col_clean = df[col].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                        unique_vals = set(col_clean.dropna().unique())
                        dict_keys_set = set(mapping.keys())
                        
                        if len(unique_vals.intersection(dict_keys_set)) > 0:
                            meaning_col = f"{col}_含义"
                            if meaning_col not in df.columns:
                                df[meaning_col] = col_clean.map(mapping).fillna(df[col])
                                translated_cols += 1
                            break 

            # 4. 导出纯净 CSV
            base_name = os.path.splitext(file_name)[0] 
            output_path = os.path.join(output_dir, f"翻译_{base_name}.csv")
            df.to_csv(output_path, index=False, encoding="utf-8-sig")
            
            print(f"  [完成] {file_name} -> 翻译_{base_name}.csv " + 
                  (f"(翻译了 {translated_cols} 个代号列)" if translated_cols > 0 else ""))

        except Exception as e:
            print(f"  [失败] 处理 {file_name} 报错: {e}")

if __name__ == "__main__":
    print("="*60)
    print("🚀 终极翻译引擎：全兼容 Word + 强力纠错版")
    print("="*60)

    global_dict, column_map = parse_word_metadata(DOC_DIR)
    global_dict = enrich_dict_from_data(DATA_DIR, global_dict, column_map)
    translate_data_files(DATA_DIR, OUTPUT_DIR, global_dict, column_map)
    print("\n✅ 数据翻译与学号统一彻底收官！")