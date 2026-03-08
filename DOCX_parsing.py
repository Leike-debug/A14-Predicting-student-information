from docx import Document
import re

def parse_comma_separated_docx(docx_path):
    doc = Document(docx_path)
    mapping = {}
    all_lines = []
    
    # 1. 提取所有段落并打散成行
    for para in doc.paragraphs:
        if para.text.strip():
            all_lines.extend(para.text.splitlines())
            
    # 2. 提取表格中的内容
    for table in doc.tables:
        for row in table.rows:
            row_text = ",".join([cell.text.strip() for cell in row.cells])
            all_lines.append(row_text)

    # 3. 逐行解析提取数据
    for line in all_lines:
        text = line.strip()
        
        # 跳过无用行
        if not text or "参数名" in text or "描述" in text or text.startswith("The following"):
            continue
            
        # 情况A：处理标准逗号分隔的表头字典 
        if ',' in text:
            parts = text.split(',')
            if len(parts) >= 2:
                key = parts[0].strip()
                if '"' in text:
                    split_by_quote = text.split('"')
                    value = split_by_quote[1].strip() if len(split_by_quote) >= 3 else parts[-1].strip()
                else:
                    value = parts[-1].strip()
                if key:
                    mapping[key] = value
            continue
            
        # 情况B：处理底部代号字典 
        # 正则表达式解释：匹配以一串数字开头，中间可能有冒号或空格，后面跟着汉字或字母
        match = re.match(r'^(\d+)[\s:：]*(.+)$', text)
        if match:
            key, value = match.groups()
            mapping[key.strip()] = value.strip()
                
    return mapping